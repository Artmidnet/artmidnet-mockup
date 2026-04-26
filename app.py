"""
Artmidnet Mockup Server — app.py V33
------------------------------------
V1:  Basic mockup generation (stretch + adapt modes)
V2:  CORS support, health check endpoint
V3:  Added /layers-report endpoint — generates DOCX, returns file directly
V4:  /layers-report now returns base64 JSON instead of file download
V5:  Page name displayed prominently as main title in report header
V6:  Added /cms-report endpoint — generates CMS Schema DOCX, returns base64 JSON
V7:  MERGE — restored /noframe, /zoom, /rect from V2 (were missing in V6)
V8:  Fixed apply_adapt — shadow paste array shape mismatch (broadcast error)
V9:  Fixed apply_zoom — use detect_outer_frame+detect_inner_canvas instead of detect_white_area
V10: Fixed apply_zoom — clip painting to inner canvas bounds (horizontal overflow fix)
V11: Fixed apply_zoom — constrain painting size to inner canvas (no overflow in any orientation)
V12: New approach for zoom+rect — detect frame, create white canvas with painting AR, add border+shadow
V13: Test — only steps A+B+C (frame detection + white canvas), no painting/border/shadow
V14: Fix detect_outer_frame — scan single pixel at center column/row instead of full row/col mean
V15: Fix detect_outer_frame — use narrow center strip (5%) mean instead of single pixel, more robust
V16: New detect_outer_frame — find largest bright canvas region, expand to cover black border
V17: New approach — detect red dot (ImagePoint) in mockup, place white canvas using size_px + painting AR
V18: שלב D — הלבשת תמונת המקור על המשטח הלבן
V19: שלב E — מסגרת שחורה בעובי 2% מממוצע גובה+רוחב
V20: שלב F — צללית רכה לימין-מטה (אור מלמעלה-שמאל)
V21: פרמטרי מסגרת וצל מגיעים מהבקשה (frame_width, frame_color, shadow_*) — ערכי default אם חסרים
V22: Fix — size_px מוגבל ל-80% מהממד הקטן של ה-mockup — מונע גלישה מחוץ לתמונה
V23: Added /receipt endpoint — builds HTML receipt and sends via Gmail SMTP (fire and forget)
V24: Fixed receipt HTML — fully inline styles, table-based layout, proper RTL for Gmail
V25: Receipt — light header bg, receipt number centered+large, fixed totals/payment direction, translate "None"
V26: Receipt email — attach PDF (weasyprint) + HTML in body
V27: Replace weasyprint with xhtml2pdf — no system dependencies required
V28: Replace xhtml2pdf with fpdf2 — pure Python, no system dependencies, Hebrew TTF font
V29: Fix RTL — only reverse Hebrew text, English painting names and Artmidnet stay as-is
V30: Redesign PDF — python-bidi for proper BiDi, beige header with logo, matches HTML email

Endpoints:
  GET  /health          — health check
  POST /mockup          — generates room mockup image (stretch / adapt)
  POST /noframe         — BuildMockupNoframe — painting centered on light background
  POST /zoom            — BuildMockupZoom    — painting fills 80% of mockup
  POST /rect            — BuildMockupRect    — adapts frame AR to painting
  POST /layers-report   — generates Layers Schema DOCX, returns base64 JSON
  POST /cms-report      — generates CMS Schema DOCX, returns base64 JSON
  POST /receipt         — builds HTML receipt and sends email to customer
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import requests
import numpy as np
from PIL import Image, ImageDraw, ImageFilter
import io
import base64
import datetime
import os
import smtplib
import threading
import sys
import types
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# V30: inject mock bidi module so fpdf2 never tries to import python-bidi
_bidi_pkg  = types.ModuleType("bidi")
_bidi_algo = types.ModuleType("bidi.algorithm")
_bidi_algo.get_display = lambda text, **kwargs: text
sys.modules.setdefault("bidi",           _bidi_pkg)
sys.modules.setdefault("bidi.algorithm", _bidi_algo)

from fpdf import FPDF

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)


# ─────────────────────────────────────────────
# Helper: download image from URL → PIL Image
# ─────────────────────────────────────────────
def load_image_from_url(url: str) -> Image.Image:
    response = requests.get(url, timeout=30)
    response.raise_for_status()
    return Image.open(io.BytesIO(response.content)).convert("RGBA")


# ─────────────────────────────────────────────
# Helper: parse hex color string → (R, G, B)
# ─────────────────────────────────────────────
def parse_hex_color(hex_str: str) -> tuple:
    """Convert '#2c1a0e' or '2c1a0e' to (R, G, B)."""
    h = hex_str.lstrip('#')
    try:
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
    except Exception:
        return (0, 0, 0)


# ─────────────────────────────────────────────
# Helper: detect red dot (ImagePoint) in mockup
# ─────────────────────────────────────────────
def detect_red_dot(img: Image.Image) -> tuple:
    arr = np.array(img.convert("RGB"))
    red_mask = (arr[:, :, 0] > 180) & (arr[:, :, 1] < 80) & (arr[:, :, 2] < 80)
    ys, xs = np.where(red_mask)
    if len(xs) == 0:
        h, w = arr.shape[:2]
        cx, cy = w // 2, h // 2
        print(f"V17 detect_red_dot: NOT FOUND — using center ({cx},{cy})")
        return cx, cy
    cx = int(np.mean(xs))
    cy = int(np.mean(ys))
    print(f"V17 detect_red_dot: found {len(xs)} red pixels | center=({cx},{cy})")
    return cx, cy


# ─────────────────────────────────────────────
# Helper: detect inner canvas (bright area)
# ─────────────────────────────────────────────
def detect_inner_canvas(arr: np.ndarray, outer: tuple, bright_threshold: int = 100) -> tuple:
    left_o, top_o, right_o, bottom_o = outer

    inner_top = top_o
    for y in range(top_o, bottom_o):
        if np.mean(arr[y, left_o:right_o]) > bright_threshold:
            inner_top = y
            break

    inner_bottom = bottom_o
    for y in range(bottom_o, top_o, -1):
        if np.mean(arr[y, left_o:right_o]) > bright_threshold:
            inner_bottom = y
            break

    inner_left = left_o
    for x in range(left_o, right_o):
        if np.mean(arr[top_o:bottom_o, x]) > bright_threshold:
            inner_left = x
            break

    inner_right = right_o
    for x in range(right_o, left_o, -1):
        if np.mean(arr[top_o:bottom_o, x]) > bright_threshold:
            inner_right = x
            break

    return inner_left, inner_top, inner_right, inner_bottom


# ─────────────────────────────────────────────
# Helper: detect white area in mockup image
# ─────────────────────────────────────────────
def detect_white_area(img: Image.Image, white_threshold: int = 240) -> tuple:
    arr = np.array(img.convert("RGB"))
    h, w = arr.shape[:2]

    white_mask = np.all(arr > white_threshold, axis=2)

    rows = np.any(white_mask, axis=1)
    cols = np.any(white_mask, axis=0)

    row_indices = np.where(rows)[0]
    col_indices = np.where(cols)[0]

    if len(row_indices) == 0 or len(col_indices) == 0:
        return w // 4, h // 4, 3 * w // 4, 3 * h // 4

    return int(col_indices[0]), int(row_indices[0]), int(col_indices[-1]), int(row_indices[-1])


# ─────────────────────────────────────────────
# Helper: extract shadow strips
# ─────────────────────────────────────────────
def extract_shadow(img_arr: np.ndarray, outer: tuple, thickness: int = 8) -> dict:
    l, t, r, b = outer
    return {
        "left":   img_arr[t:b + 1, l:l + thickness].copy(),
        "bottom": img_arr[b - thickness:b + 1, l:r + 1].copy(),
    }


# ─────────────────────────────────────────────
# Helper: sample wall color near frame
# ─────────────────────────────────────────────
def sample_wall_color(img_arr: np.ndarray, outer: tuple, margin: int = 20) -> tuple:
    l, t, r, b = outer
    y0 = max(0, t - margin)
    y1 = t
    x0 = l + (r - l) // 3
    x1 = r - (r - l) // 3
    patch = img_arr[y0:y1, x0:x1]
    if patch.size == 0:
        return (200, 200, 200, 255)
    mean = patch.mean(axis=(0, 1))
    return tuple(int(v) for v in mean[:4])


# ─────────────────────────────────────────────
# Helper: sample background color from corners
# ─────────────────────────────────────────────
def sample_corner_color(img: Image.Image, corner_size: int = 60) -> tuple:
    arr = np.array(img.convert("RGBA"))
    h, w = arr.shape[:2]
    cs = min(corner_size, h // 4, w // 4)
    corners = [
        arr[:cs, :cs],
        arr[:cs, w - cs:],
        arr[h - cs:, :cs],
        arr[h - cs:, w - cs:]
    ]
    mean = np.mean([c.mean(axis=(0, 1)) for c in corners], axis=0)
    return tuple(int(v) for v in mean)


# ─────────────────────────────────────────────
# Helper: PIL Image → base64 JPEG
# ─────────────────────────────────────────────
def image_to_base64(img: Image.Image, fmt: str = "JPEG") -> str:
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format=fmt, quality=90)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


# ─────────────────────────────────────────────
# MODE: STRETCH
# ─────────────────────────────────────────────
def apply_stretch(room_img: Image.Image, painting_img: Image.Image) -> Image.Image:
    arr = np.array(room_img)
    outer = detect_outer_frame(room_img)
    inner = detect_inner_canvas(arr, outer)

    il, it, ir, ib = inner
    canvas_w = ir - il
    canvas_h = ib - it

    if canvas_w <= 0 or canvas_h <= 0:
        raise ValueError("Could not detect inner canvas area")

    painting_resized = painting_img.resize((canvas_w, canvas_h), Image.LANCZOS).convert("RGBA")
    result = room_img.copy().convert("RGBA")
    result.paste(painting_resized, (il, it), painting_resized)
    return result.convert("RGB")


# ─────────────────────────────────────────────
# MODE: ADAPT
# ─────────────────────────────────────────────
def apply_adapt(room_img: Image.Image, painting_img: Image.Image) -> Image.Image:
    arr = np.array(room_img.convert("RGBA"))
    outer = detect_outer_frame(room_img)
    inner = detect_inner_canvas(arr, outer)

    lo, to, ro, bo = outer
    li, ti, ri, bi = inner

    ft_left   = li - lo
    ft_top    = ti - to
    ft_right  = ro - ri
    ft_bottom = bo - bi

    orig_canvas_w = ri - li
    paint_w, paint_h = painting_img.size
    aspect = paint_h / paint_w
    new_canvas_w = orig_canvas_w
    new_canvas_h = int(new_canvas_w * aspect)

    new_frame_w = new_canvas_w + ft_left + ft_right
    new_frame_h = new_canvas_h + ft_top  + ft_bottom

    shadow = extract_shadow(arr, outer)
    wall_color = sample_wall_color(arr, outer)

    frame_arr = np.zeros((new_frame_h, new_frame_w, 4), dtype=np.uint8)
    frame_arr[:, :] = (0, 0, 0, 255)

    painting_resized = painting_img.resize((new_canvas_w, new_canvas_h), Image.LANCZOS).convert("RGBA")
    p_arr = np.array(painting_resized)
    frame_arr[ft_top:ft_top + new_canvas_h, ft_left:ft_left + new_canvas_w] = p_arr

    frame_img = Image.fromarray(frame_arr, "RGBA")
    result = room_img.copy().convert("RGBA")

    result_arr = np.array(result)
    img_h, img_w = result_arr.shape[:2]

    result_arr[to:bo + 1, lo:ro + 1] = wall_color
    result = Image.fromarray(result_arr, "RGBA")
    result.paste(frame_img, (lo, to), frame_img)

    result_arr2 = np.array(result)

    sh_left = shadow["left"]
    if sh_left.size > 0:
        new_sh_left = np.array(Image.fromarray(sh_left).resize(
            (sh_left.shape[1], new_frame_h), Image.LANCZOS))
        x0 = lo - sh_left.shape[1]
        if x0 >= 0:
            y_end = min(to + new_frame_h, img_h)
            actual_h = y_end - to
            result_arr2[to:y_end, x0:lo] = new_sh_left[:actual_h, :]

    sh_bot = shadow["bottom"]
    if sh_bot.size > 0:
        new_sh_bot = np.array(Image.fromarray(sh_bot).resize(
            (new_frame_w, sh_bot.shape[0]), Image.LANCZOS))
        y0 = to + new_frame_h
        if y0 < img_h:
            y_end = min(y0 + sh_bot.shape[0], img_h)
            actual_h = y_end - y0
            x_end = min(lo + new_frame_w, img_w)
            actual_w = x_end - lo
            result_arr2[y0:y_end, lo:x_end] = new_sh_bot[:actual_h, :actual_w]

    return Image.fromarray(result_arr2, "RGBA").convert("RGB")


# ─────────────────────────────────────────────
# MODE: NOFRAME
# ─────────────────────────────────────────────
def apply_noframe(painting_img: Image.Image) -> Image.Image:
    pw, ph = painting_img.size
    if max(pw, ph) != 2000:
        scale = 2000 / max(pw, ph)
        pw = int(pw * scale)
        ph = int(ph * scale)
        painting_img = painting_img.resize((pw, ph), Image.LANCZOS)

    arr = np.array(painting_img.convert("RGB")).reshape(-1, 3).astype(float)
    indices = np.random.choice(len(arr), min(300, len(arr)), replace=False)
    chosen = arr[indices[np.random.randint(len(indices))]]
    light = tuple(int(230 + (c / 255.0) * 20) for c in chosen) + (255,)

    canvas = Image.new("RGBA", (2000, 2000), light)

    x = (2000 - pw) // 2
    y = (2000 - ph) // 2

    painting_rgba = painting_img.convert("RGBA")
    canvas.paste(painting_rgba, (x, y), painting_rgba)

    return canvas.convert("RGB")


# ─────────────────────────────────────────────
# MODE: apply_new_mockup (zoom + rect)
# ─────────────────────────────────────────────
def apply_new_mockup(
    painting_img: Image.Image,
    mockup_img: Image.Image,
    size_px: int = 800,
    frame_width: float = None,
    frame_color: str = None,
    shadow_offset_x: float = None,
    shadow_offset_y: float = None,
    shadow_blur: float = None,
    shadow_spread: float = None,
    shadow_opacity: float = None
) -> Image.Image:

    img_cx, img_cy = detect_red_dot(mockup_img)

    img_w, img_h = mockup_img.size
    max_size = int(min(img_w, img_h) * 0.80)
    if size_px > max_size:
        print(f"V22 size_px capped: {size_px} → {max_size} (mockup={img_w}x{img_h})")
        size_px = max_size

    pw, ph = painting_img.size
    ar = ph / pw
    print(f"V22 painting: {pw}x{ph} AR={ar:.3f} | size_px={size_px}")

    if ar <= 1.0:
        wc_w = size_px
        wc_h = int(size_px * ar)
    else:
        wc_h = size_px
        wc_w = int(size_px / ar)

    print(f"V22 white canvas: {wc_w}x{wc_h} | ImagePoint=({img_cx},{img_cy})")

    avg_side = (wc_w + wc_h) / 2
    _frame_color      = parse_hex_color(frame_color) if frame_color else (0, 0, 0)
    _border_thickness = max(1, round(frame_width)) if frame_width is not None else max(1, round(avg_side * 0.02))
    _shadow_offset_x  = int(shadow_offset_x) if shadow_offset_x is not None else max(2, round(wc_w * 0.03))
    _shadow_offset_y  = int(shadow_offset_y) if shadow_offset_y is not None else max(2, round(wc_h * 0.03))
    _shadow_blur      = int(shadow_blur)      if shadow_blur      is not None else max(3, round(avg_side * 0.025))
    _shadow_opacity   = int(shadow_opacity * 255) if shadow_opacity is not None else 100

    painting_resized = painting_img.convert("RGBA").resize((wc_w, wc_h), Image.LANCZOS)

    paste_x = img_cx - wc_w // 2
    paste_y = img_cy - wc_h // 2
    paste_x = max(0, min(paste_x, img_w - wc_w))
    paste_y = max(0, min(paste_y, img_h - wc_h))

    shadow_layer = Image.new("RGBA", mockup_img.size, (0, 0, 0, 0))
    shadow_draw  = ImageDraw.Draw(shadow_layer)
    sx = paste_x + _shadow_offset_x
    sy = paste_y + _shadow_offset_y
    shadow_draw.rectangle(
        [sx, sy, sx + wc_w - 1, sy + wc_h - 1],
        fill=(0, 0, 0, _shadow_opacity)
    )
    shadow_layer = shadow_layer.filter(ImageFilter.GaussianBlur(radius=_shadow_blur))

    base = mockup_img.copy().convert("RGBA")
    base.paste(shadow_layer, (0, 0), shadow_layer)
    base.paste(painting_resized, (paste_x, paste_y), painting_resized)

    draw = ImageDraw.Draw(base)
    for i in range(_border_thickness):
        draw.rectangle(
            [paste_x + i, paste_y + i,
             paste_x + wc_w - 1 - i, paste_y + wc_h - 1 - i],
            outline=_frame_color + (255,)
        )

    print(f"V22 done: paste=({paste_x},{paste_y}) size={wc_w}x{wc_h}")
    return base.convert("RGB")


# ─────────────────────────────────────────────
# MODE: ZOOM
# ─────────────────────────────────────────────
def apply_zoom(painting_img: Image.Image, mockup_img: Image.Image, size_px: int = 800, **kwargs) -> Image.Image:
    return apply_new_mockup(painting_img, mockup_img, size_px, **kwargs)


# ─────────────────────────────────────────────
# MODE: RECT
# ─────────────────────────────────────────────
def apply_rect(painting_img: Image.Image, mockup_img: Image.Image, size_px: int = 800, **kwargs) -> Image.Image:
    return apply_new_mockup(painting_img, mockup_img, size_px, **kwargs)


# ═════════════════════════════════════════════
# RECEIPT: HTML Builder
# ═════════════════════════════════════════════

def build_receipt_html(data: dict) -> str:
    """V25: RTL receipt — light header, centered receipt number, fixed directions, translate None."""

    # ── colors ──
    C_DARK    = "#1a2e4a"
    C_GOLD    = "#c9a84c"
    C_WHITE   = "#ffffff"
    C_HEADER  = "#f8f4ef"   # V25: בהיר מאוד לheader
    C_BG      = "#f5f5f5"
    C_BORDER  = "#e0e0e0"
    C_TEXT    = "#333333"
    C_MUTED   = "#888888"
    C_LIGHT   = "#fafafa"

    # ── VAT ──
    vat_rate  = data.get("vatRate", 0)
    vat_label = f"מע\"מ {int(vat_rate * 100)}%" if (vat_rate and vat_rate > 0) else "מע\"מ (פטור)"
    vat_value = data.get("vatAmount", "₪0.00") if (vat_rate and vat_rate > 0) else "פטור"

    # ── logo ──
    logo_url = data.get("logoUrl", "")
    if logo_url:
        logo_html = f'<img src="{logo_url}" alt="לוגו" width="130" style="display:block;max-height:60px;width:auto;">'
    else:
        logo_html = f'<span style="font-size:18px;font-weight:bold;color:{C_DARK};">{data.get("businessName","")}</span>'

    # ── items rows — V25: translate "None" → "ללא מסגרת" ──
    def translate_details(details: str) -> str:
        return details.replace("None", "ללא מסגרת").replace("none", "ללא מסגרת")

    items_html = ""
    for i, item in enumerate(data.get("items", [])):
        row_bg  = C_LIGHT if i % 2 == 0 else C_WHITE
        details = translate_details(item.get('details', ''))
        items_html += f"""
        <tr>
          <td style="padding:10px 8px;border-bottom:1px solid {C_BORDER};text-align:right;font-size:12px;color:{C_TEXT};background:{row_bg};">{item.get('index','')}</td>
          <td style="padding:10px 8px;border-bottom:1px solid {C_BORDER};text-align:right;font-size:12px;background:{row_bg};">
            <div style="font-weight:bold;color:{C_DARK};">{item.get('name','')}</div>
            <div style="font-size:11px;color:{C_MUTED};margin-top:2px;">{details}</div>
          </td>
          <td style="padding:10px 8px;border-bottom:1px solid {C_BORDER};text-align:center;font-size:12px;color:{C_TEXT};background:{row_bg};">1</td>
          <td style="padding:10px 8px;border-bottom:1px solid {C_BORDER};text-align:right;font-size:12px;color:{C_TEXT};background:{row_bg};">{item.get('price','')}</td>
          <td style="padding:10px 8px;border-bottom:1px solid {C_BORDER};text-align:right;font-size:12px;font-weight:bold;color:{C_DARK};background:{row_bg};">{item.get('total','')}</td>
        </tr>"""

    # ── payment detail row ──
    payment_details     = data.get("paymentDetails", "")
    payment_detail_row  = ""
    if payment_details:
        payment_detail_row = f"""
        <tr>
          <td style="padding:8px 12px;text-align:right;font-size:12px;color:{C_MUTED};border-bottom:1px solid {C_BORDER};width:50%;">פירוט</td>
          <td style="padding:8px 12px;text-align:right;font-size:12px;font-weight:bold;color:{C_TEXT};border-bottom:1px solid {C_BORDER};">{payment_details}</td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html lang="he" dir="rtl">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>{data.get('documentType','קבלה')} {data.get('receiptNumber','')}</title>
</head>
<body style="margin:0;padding:20px;background-color:#e8e8e8;font-family:Arial,sans-serif;direction:rtl;">

<table width="100%" cellpadding="0" cellspacing="0" border="0" style="max-width:680px;margin:0 auto;">
<tr><td>

  <!-- DOCUMENT WRAPPER -->
  <table width="100%" cellpadding="0" cellspacing="0" border="0"
    style="background:{C_WHITE};border:1px solid {C_BORDER};border-radius:4px;">

    <!-- HEADER — V25: light bg -->
    <tr>
      <td style="background:{C_HEADER};padding:20px 28px;border-radius:4px 4px 0 0;border-bottom:1px solid {C_BORDER};">
        <table width="100%" cellpadding="0" cellspacing="0" border="0">
          <tr>
            <td style="text-align:right;vertical-align:middle;">{logo_html}</td>
            <td style="text-align:left;vertical-align:middle;">
              <div style="font-size:15px;font-weight:bold;color:{C_DARK};margin-bottom:4px;">{data.get('businessName','')}</div>
              <div style="font-size:11px;color:#555;line-height:1.8;">
                ח.פ. {data.get('businessTaxId','')}<br>
                {data.get('bizAddress','')}<br>
                {data.get('businessEmail','')} | {data.get('bizPhone','')}
              </div>
            </td>
          </tr>
        </table>
      </td>
    </tr>

    <!-- TITLE BAND — V25: receipt number centered + large -->
    <tr>
      <td style="background:{C_GOLD};padding:14px 28px;">
        <table width="100%" cellpadding="0" cellspacing="0" border="0">
          <tr>
            <!-- שורה 1: מספר קבלה מרוכז -->
            <td colspan="2" style="text-align:center;padding-bottom:6px;">
              <span style="font-size:26px;font-weight:bold;color:{C_DARK};">
                {data.get('documentType','קבלה')} מספר {data.get('receiptNumber','')}
              </span>
            </td>
          </tr>
          <tr>
            <!-- שורה 2: מספר הזמנה קטן יותר -->
            <td colspan="2" style="text-align:center;">
              <span style="font-size:12px;color:{C_DARK};opacity:0.8;">
                הזמנה מספר {data.get('orderNumber','')}
              </span>
            </td>
          </tr>
        </table>
      </td>
    </tr>

    <!-- BODY -->
    <tr>
      <td style="padding:24px 28px;background:{C_WHITE};">

        <!-- RECIPIENT + DATE -->
        <table width="100%" cellpadding="0" cellspacing="0" border="0"
          style="margin-bottom:20px;padding-bottom:16px;border-bottom:1px solid {C_BORDER};">
          <tr>
            <td style="text-align:right;vertical-align:top;">
              <div style="font-size:10px;font-weight:bold;color:{C_MUTED};text-transform:uppercase;letter-spacing:1px;margin-bottom:5px;">לכבוד</div>
              <div style="font-size:15px;font-weight:bold;color:{C_DARK};">{data.get('customerName','')}</div>
              <div style="font-size:12px;color:#666;margin-top:3px;">{data.get('customerPhone','')}</div>
              <div style="font-size:12px;color:#666;margin-top:2px;">{data.get('customerEmail','')}</div>
            </td>
            <td style="text-align:left;vertical-align:top;">
              <div style="font-size:10px;font-weight:bold;color:{C_MUTED};text-transform:uppercase;letter-spacing:1px;margin-bottom:5px;">תאריך</div>
              <div style="font-size:14px;font-weight:bold;color:{C_DARK};">{data.get('orderDate','')}</div>
            </td>
          </tr>
        </table>

        <!-- ITEMS SECTION TITLE -->
        <div style="font-size:10px;font-weight:bold;color:{C_MUTED};text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;text-align:right;">פירוט הרכישה</div>

        <!-- ITEMS TABLE -->
        <table width="100%" cellpadding="0" cellspacing="0" border="0"
          style="border-collapse:collapse;margin-bottom:20px;font-size:12px;">
          <thead>
            <tr style="background:{C_DARK};">
              <th style="padding:9px 8px;text-align:right;font-size:11px;font-weight:bold;color:{C_WHITE};width:36px;">מק"ט</th>
              <th style="padding:9px 8px;text-align:right;font-size:11px;font-weight:bold;color:{C_WHITE};">פירוט</th>
              <th style="padding:9px 8px;text-align:center;font-size:11px;font-weight:bold;color:{C_WHITE};width:50px;">כמות</th>
              <th style="padding:9px 8px;text-align:right;font-size:11px;font-weight:bold;color:{C_WHITE};width:80px;">מחיר</th>
              <th style="padding:9px 8px;text-align:right;font-size:11px;font-weight:bold;color:{C_WHITE};width:80px;">סה"כ</th>
            </tr>
          </thead>
          <tbody>{items_html}</tbody>
        </table>

        <!-- TOTALS — V25: fixed direction, סכום מימין -->
        <table width="100%" cellpadding="0" cellspacing="0" border="0" style="margin-bottom:20px;">
          <tr>
            <td width="50%"></td>
            <td width="50%">
              <table width="100%" cellpadding="0" cellspacing="0" border="0"
                style="border:1px solid {C_BORDER};border-radius:3px;overflow:hidden;">
                <tr>
                  <td style="padding:8px 12px;text-align:right;font-size:12px;color:#666;border-bottom:1px solid {C_BORDER};">סכום ביניים</td>
                  <td style="padding:8px 12px;text-align:left;font-size:12px;font-weight:bold;color:{C_TEXT};border-bottom:1px solid {C_BORDER};">{data.get('subtotal','')}</td>
                </tr>
                <tr>
                  <td style="padding:8px 12px;text-align:right;font-size:12px;color:#666;border-bottom:1px solid {C_BORDER};">משלוח</td>
                  <td style="padding:8px 12px;text-align:left;font-size:12px;font-weight:bold;color:{C_TEXT};border-bottom:1px solid {C_BORDER};">{data.get('shipping','₪0.00')}</td>
                </tr>
                <tr>
                  <td style="padding:8px 12px;text-align:right;font-size:11px;color:#aaa;border-bottom:1px solid {C_BORDER};">{vat_label}</td>
                  <td style="padding:8px 12px;text-align:left;font-size:11px;color:#aaa;border-bottom:1px solid {C_BORDER};">{vat_value}</td>
                </tr>
                <tr style="background:{C_DARK};">
                  <td style="padding:10px 12px;text-align:right;font-size:13px;font-weight:bold;color:{C_WHITE};">סה"כ לתשלום</td>
                  <td style="padding:10px 12px;text-align:left;font-size:15px;font-weight:bold;color:{C_GOLD};">{data.get('total','')}</td>
                </tr>
              </table>
            </td>
          </tr>
        </table>

        <!-- PAYMENT SECTION TITLE -->
        <div style="font-size:10px;font-weight:bold;color:{C_MUTED};text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;text-align:right;">פרטי תשלום</div>

        <!-- PAYMENT TABLE — V25: fixed direction -->
        <table width="100%" cellpadding="0" cellspacing="0" border="0"
          style="background:{C_BG};border:1px solid {C_BORDER};border-radius:3px;margin-bottom:20px;">
          <tr>
            <td style="padding:8px 12px;text-align:right;font-size:12px;font-weight:bold;color:{C_TEXT};border-bottom:1px solid {C_BORDER};width:50%;">{data.get('paymentMethod','')}</td>
            <td style="padding:8px 12px;text-align:left;font-size:12px;color:{C_MUTED};border-bottom:1px solid {C_BORDER};">אמצעי תשלום</td>
          </tr>
          {payment_detail_row}
          <tr>
            <td style="padding:8px 12px;text-align:right;font-size:12px;font-weight:bold;color:{C_TEXT};border-bottom:1px solid {C_BORDER};">{data.get('orderDate','')}</td>
            <td style="padding:8px 12px;text-align:left;font-size:12px;color:{C_MUTED};border-bottom:1px solid {C_BORDER};">תאריך חיוב</td>
          </tr>
          <tr>
            <td style="padding:8px 12px;text-align:right;font-size:12px;font-weight:bold;color:{C_TEXT};">{data.get('total','')}</td>
            <td style="padding:8px 12px;text-align:left;font-size:12px;color:{C_MUTED};">סכום</td>
          </tr>
        </table>

      </td>
    </tr>

    <!-- FOOTER -->
    <tr>
      <td style="background:{C_BG};border-top:2px solid {C_DARK};padding:10px 28px;border-radius:0 0 4px 4px;">
        <table width="100%" cellpadding="0" cellspacing="0" border="0">
          <tr>
            <td style="text-align:right;font-size:10px;color:{C_MUTED};">{data.get('footerText','')}</td>
            <td style="text-align:left;font-size:10px;color:#bbb;">{data.get('documentType','קבלה')} מס&apos; {data.get('receiptNumber','')} | עמוד 1 מתוך 1</td>
          </tr>
        </table>
      </td>
    </tr>

  </table>
  <!-- END DOCUMENT WRAPPER -->

</td></tr>
</table>

</body>
</html>"""

    return html


# ═════════════════════════════════════════════
# RECEIPT: PDF Builder (V30 — fpdf2 + python-bidi + new design)
# ═════════════════════════════════════════════

def build_receipt_pdf(data: dict) -> bytes:
    """V30: PDF receipt — manual RTL reversal, no python-bidi dependency."""
    import tempfile

    # ── font path ──
    font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "NotoSansHebrew-Regular.ttf")
    print(f"V33 build_receipt_pdf: font={font_path} exists={os.path.exists(font_path)}")

    # ── RTL helpers ──
    def has_hebrew(text: str) -> bool:
        return any('א' <= c <= 'ת' for c in str(text))

    def bidi(text: str) -> str:
        """Reverse text only if it contains Hebrew — fpdf2 renders LTR so we pre-reverse."""
        if not text:
            return ""
        t = str(text)
        return t[::-1] if has_hebrew(t) else t

    # ── colors ──
    C_DARK  = (26, 46, 74)    # #1a2e4a
    C_GOLD  = (201, 168, 76)  # #c9a84c
    C_HEAD  = (248, 244, 239) # #f8f4ef beige header
    C_BORD  = (224, 224, 224) # #e0e0e0
    C_MUTED = (136, 136, 136)
    C_TEXT  = (50, 50, 50)
    C_LIGHT = (250, 250, 250)

    # ── download logo ──
    logo_path = None
    logo_url  = data.get("logoUrl", "")
    if logo_url:
        try:
            r = requests.get(logo_url, timeout=10)
            r.raise_for_status()
            suffix = ".png" if "png" in logo_url.lower() else ".jpg"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(r.content)
                logo_path = tmp.name
            print(f"V30 logo downloaded: {logo_path}")
        except Exception as e:
            print(f"V30 logo download failed: {e}")

    # ── data fields ──
    business_name  = data.get("businessName", "Artmidnet")
    tax_id         = data.get("businessTaxId", "")
    biz_addr       = data.get("bizAddress", "")
    biz_email      = data.get("businessEmail", "")
    biz_phone      = data.get("bizPhone", "")
    doc_type       = data.get("documentType", "קבלה")
    receipt_num    = str(data.get("receiptNumber", ""))
    order_number   = str(data.get("orderNumber", ""))
    customer_name  = data.get("customerName", "")
    customer_email = data.get("customerEmail", "")
    customer_phone = data.get("customerPhone", "")
    order_date     = data.get("orderDate", "")
    subtotal       = str(data.get("subtotal", ""))
    shipping       = str(data.get("shipping", "₪0.00"))
    total          = str(data.get("total", ""))
    vat_rate       = data.get("vatRate", 0)
    vat_value      = str(data.get("vatAmount", "פטור")) if (vat_rate and vat_rate > 0) else "פטור"
    vat_label      = f'מע"מ {int(vat_rate * 100)}%' if (vat_rate and vat_rate > 0) else 'מע"מ (פטור)'
    payment_method  = str(data.get("paymentMethod", ""))
    payment_details = str(data.get("paymentDetails", ""))
    footer_text    = data.get("footerText", "")
    items          = data.get("items", [])

    # ── init PDF ──
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("Hebrew", fname=font_path)
    page_w = pdf.w  # 210mm A4

    # ════════════════════════════════
    # HEADER — beige bg, logo right, biz info left
    # ════════════════════════════════
    header_h = 34
    pdf.set_fill_color(*C_HEAD)
    pdf.rect(0, 0, page_w, header_h, style="F")

    # logo (right side)
    if logo_path:
        try:
            pdf.image(logo_path, x=page_w - 55, y=4, h=26)
        except Exception as e:
            print(f"V30 logo embed failed: {e}")

    # business name (left side)
    pdf.set_text_color(*C_DARK)
    pdf.set_font("Hebrew", size=13)
    pdf.set_xy(10, 5)
    pdf.cell(140, 8, business_name)

    # business details (left side)
    pdf.set_font("Hebrew", size=8)
    pdf.set_text_color(85, 85, 85)
    pdf.set_xy(10, 14)
    pdf.cell(140, 5, f"ח.פ. {tax_id}")
    pdf.set_xy(10, 19)
    pdf.cell(140, 5, bidi(biz_addr))
    pdf.set_xy(10, 24)
    pdf.cell(140, 5, f"{biz_email} | {biz_phone}")

    # bottom border of header
    pdf.set_draw_color(*C_BORD)
    pdf.set_line_width(0.3)
    pdf.line(0, header_h, page_w, header_h)

    # ════════════════════════════════
    # GOLD BAND — receipt number centered
    # ════════════════════════════════
    gold_y = header_h
    gold_h = 24
    pdf.set_fill_color(*C_GOLD)
    pdf.rect(0, gold_y, page_w, gold_h, style="F")

    pdf.set_text_color(*C_DARK)
    pdf.set_font("Hebrew", size=18)
    pdf.set_xy(10, gold_y + 2)
    pdf.cell(page_w - 20, 12, bidi(f"{doc_type} מספר {receipt_num}"), align="C")

    pdf.set_font("Hebrew", size=9)
    pdf.set_xy(10, gold_y + 14)
    pdf.cell(page_w - 20, 8, bidi(f"הזמנה מספר {order_number}"), align="C")

    # ════════════════════════════════
    # BODY
    # ════════════════════════════════
    body_y = gold_y + gold_h + 5
    pdf.set_y(body_y)

    # ── customer + date ──
    pdf.set_font("Hebrew", size=8)
    pdf.set_text_color(*C_MUTED)
    pdf.set_xy(10, body_y)
    pdf.cell(95, 5, "תאריך", align="L")
    pdf.cell(95, 5, bidi("לכבוד"), align="R", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Hebrew", size=12)
    pdf.set_text_color(*C_DARK)
    pdf.set_x(10)
    pdf.cell(95, 7, order_date, align="L")
    pdf.cell(95, 7, bidi(customer_name), align="R", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Hebrew", size=9)
    pdf.set_text_color(100, 100, 100)
    pdf.set_x(10)
    pdf.cell(95, 5, customer_email, align="L")
    pdf.cell(95, 5, customer_phone, align="R", new_x="LMARGIN", new_y="NEXT")

    # ── separator ──
    sep_y = pdf.get_y() + 3
    pdf.set_draw_color(*C_BORD)
    pdf.set_line_width(0.3)
    pdf.line(10, sep_y, page_w - 10, sep_y)
    pdf.set_y(sep_y + 4)

    # ── section label ──
    pdf.set_font("Hebrew", size=8)
    pdf.set_text_color(*C_MUTED)
    pdf.set_x(10)
    pdf.cell(page_w - 20, 5, bidi('פירוט הרכישה'), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)

    # ── items table header — V32: reversed column order for RTL layout ──
    # visual RTL order (left→right on page): סה"כ | מחיר | כמות | פירוט | מק"ט
    col_w   = [30, 30, 20, 95, 15]
    headers = ['סה"כ', "מחיר", "כמות", "פירוט", 'מק"ט']
    pdf.set_fill_color(*C_DARK)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Hebrew", size=9)
    pdf.set_x(10)
    for w, h in zip(col_w, headers):
        pdf.cell(w, 7, bidi(h), border=0, align="C", fill=True)
    pdf.ln()

    # ── items rows ──
    for idx, item in enumerate(items):
        even = idx % 2 == 0
        pdf.set_fill_color(*C_LIGHT) if even else pdf.set_fill_color(255, 255, 255)
        pdf.set_text_color(*C_TEXT)
        pdf.set_font("Hebrew", size=8)
        pdf.set_x(10)

        item_index = str(item.get("index", ""))
        item_name  = item.get("name", "")
        item_price = str(item.get("price", ""))
        item_total = str(item.get("total", ""))
        details    = item.get("details", "").replace("None", "ללא מסגרת").replace("none", "ללא מסגרת")

        # V32: reverse PART ORDER only — PDF viewer bidi handles character display
        parts = [p.strip() for p in details.split("|")] if details else []
        details_rtl = " | ".join(parts) if parts else ""  # V33: natural order, PDF viewer renders RTL

        # V32: reversed column order matches header (סה"כ left, מק"ט right)
        pdf.cell(col_w[0], 6, item_total,   border=0,   align="C",  fill=even)
        pdf.cell(col_w[1], 6, item_price,   border=0,   align="C",  fill=even)
        pdf.cell(col_w[2], 6, "1",           border=0,   align="C",  fill=even)
        pdf.cell(col_w[3], 6, item_name,    border=0,   align="R",  fill=even)
        pdf.cell(col_w[4], 6, item_index,   border=0,   align="C",  fill=even)
        pdf.ln()
        # second line — Hebrew details right-aligned in the wide פירוט column
        if details_rtl:
            pdf.set_font("Hebrew", size=7)
            pdf.set_text_color(*C_MUTED)
            pdf.set_x(10)
            pdf.cell(col_w[0], 5, "",           border="B", align="C", fill=even)
            pdf.cell(col_w[1], 5, "",           border="B", align="C", fill=even)
            pdf.cell(col_w[2], 5, "",           border="B", align="C", fill=even)
            pdf.cell(col_w[3], 5, details_rtl,  border="B", align="R", fill=even)
            pdf.cell(col_w[4], 5, "",           border="B", align="C", fill=even)
            pdf.ln()
            pdf.set_font("Hebrew", size=8)
            pdf.set_text_color(*C_TEXT)
        else:
            pdf.set_x(10)
            for w in col_w:
                pdf.cell(w, 1, "", border="B", fill=even)
            pdf.ln()

    # ── totals ──
    pdf.ln(4)
    pdf.set_text_color(*C_TEXT)

    def totals_row(label, value, font_size=10, fill=False, bg=None, fg=None):
        if bg:
            pdf.set_fill_color(*bg)
        if fg:
            pdf.set_text_color(*fg)
        pdf.set_font("Hebrew", size=font_size)
        pdf.set_x(10)
        pdf.cell(95, 8, str(value), align="L", fill=fill)
        pdf.cell(95, 8, bidi(label), align="R", fill=fill, new_x="LMARGIN", new_y="NEXT")

    totals_row("סכום ביניים", subtotal)
    totals_row("משלוח", shipping)
    pdf.set_text_color(170, 170, 170)
    totals_row(vat_label, bidi(vat_value), font_size=9)
    totals_row('סה"כ לתשלום', total, font_size=13, fill=True, bg=C_DARK, fg=C_GOLD)

    # ── payment ──
    pdf.ln(4)
    pdf.set_text_color(*C_MUTED)
    pdf.set_font("Hebrew", size=8)
    pdf.set_x(10)
    pdf.cell(page_w - 20, 5, bidi("פרטי תשלום"), align="R", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)

    def payment_row(label, value):
        pdf.set_text_color(*C_TEXT)
        pdf.set_font("Hebrew", size=9)
        pdf.set_x(10)
        # value: use bidi only if Hebrew, otherwise show as-is (e.g. PayPal, credit card)
        val_str = str(value)
        pdf.cell(95, 7, bidi(val_str) if has_hebrew(val_str) else val_str, align="L")
        pdf.cell(95, 7, bidi(label), align="R", new_x="LMARGIN", new_y="NEXT")

    payment_row("אמצעי תשלום", payment_method)
    if payment_details:
        payment_row("פירוט", payment_details)
    payment_row("תאריך חיוב", order_date)
    payment_row("סכום", total)

    # ── footer ──
    pdf.set_y(-18)
    pdf.set_draw_color(*C_DARK)
    pdf.set_line_width(0.5)
    pdf.line(10, pdf.get_y(), page_w - 10, pdf.get_y())
    pdf.ln(2)
    pdf.set_text_color(170, 170, 170)
    pdf.set_font("Hebrew", size=8)
    pdf.set_x(10)
    pdf.cell(page_w - 20, 5, bidi(footer_text), align="R")

    # ── cleanup logo temp file ──
    if logo_path and os.path.exists(logo_path):
        try:
            os.unlink(logo_path)
        except Exception:
            pass

    print(f"V33 build_receipt_pdf: PDF built successfully")
    return pdf.output()


# ─────────────────────────────────────────────
# RECEIPT: Gmail Sender (runs in background thread)
# ─────────────────────────────────────────────

def send_receipt_email(to_email: str, subject: str, html_body: str, data: dict = None):
    """V28: Send HTML receipt email via Gmail SMTP with PDF attachment (fpdf2)."""
    gmail_user = os.environ.get("GMAIL_USER", "")
    gmail_pass = os.environ.get("GMAIL_APP_PASS", "")

    if not gmail_user or not gmail_pass:
        print("V28 send_receipt_email: ERROR — GMAIL_USER or GMAIL_APP_PASS not set")
        return

    try:
        # ── Generate PDF using fpdf2 ──
        pdf_bytes    = None
        receipt_num  = str(data.get("receiptNumber", "")) if data else ""
        pdf_filename = f"receipt_{receipt_num}.pdf" if receipt_num else "receipt.pdf"

        if data:
            print("V28 send_receipt_email: generating PDF with fpdf2...")
            pdf_bytes = build_receipt_pdf(data)
            print(f"V28 send_receipt_email: PDF generated — {len(pdf_bytes)} bytes")
        else:
            print("V28 send_receipt_email: no data provided — skipping PDF")

        # ── Build email ──
        msg = MIMEMultipart("mixed")
        msg["Subject"] = subject
        msg["From"]    = gmail_user
        msg["To"]      = to_email

        # HTML body
        msg.attach(MIMEText(html_body, "html", "utf-8"))

        # PDF attachment (if generated successfully)
        if pdf_bytes:
            pdf_part = MIMEApplication(pdf_bytes, _subtype="pdf")
            pdf_part.add_header("Content-Disposition", "attachment", filename=pdf_filename)
            msg.attach(pdf_part)

        # ── Send ──
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(gmail_user, gmail_pass)
            server.sendmail(gmail_user, to_email, msg.as_string())

        print(f"V28 send_receipt_email: sent to {to_email} with PDF {pdf_filename}")

    except Exception as e:
        print(f"V28 send_receipt_email: FAILED — {str(e)}")


# ═════════════════════════════════════════════
# DOCX Helpers
# ═════════════════════════════════════════════

TYPE_COLORS = {
    'Section': 'C00000', 'Box': '2E75B6', 'Text': '375623',
    'Button': '7030A0', 'Repeater': 'C55A11', 'Image': '006400',
    'VectorImage': '006400', 'Menu': '595959', 'MenuContainer': '595959',
}

FIELD_TYPE_COLORS = {
    'TEXT':          '1F3864',
    'NUMBER':        '375623',
    'BOOLEAN':       'C55A11',
    'DATE':          '7030A0',
    'IMAGE':         '006400',
    'MEDIA_GALLERY': '006400',
    'REFERENCE':     'C00000',
    'ARRAY_STRING':  '2E75B6',
    'ARRAY':         '2E75B6',
    'OBJECT':        '595959',
    'RICH_TEXT':     '404040',
}

def get_type_color(element_type):
    return TYPE_COLORS.get(element_type, '404040')

def get_field_type_color(field_type):
    return FIELD_TYPE_COLORS.get(field_type, '404040')

def hex_to_rgb(hex_str):
    h = hex_str.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


# ═════════════════════════════════════════════
# ENDPOINTS: IMAGE MOCKUPS
# ═════════════════════════════════════════════

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok", "service": "artmidnet-mockup", "version": "V33"})


@app.route("/mockup", methods=["POST"])
def mockup():
    try:
        data = request.get_json(force=True)
        room_url     = data.get("room_url")
        painting_url = data.get("painting_url")
        mode         = data.get("mode", "stretch").lower()

        if not room_url or not painting_url:
            return jsonify({"error": "room_url and painting_url are required"}), 400
        if mode not in ("stretch", "adapt"):
            return jsonify({"error": "mode must be 'stretch' or 'adapt'"}), 400

        room_img     = load_image_from_url(room_url)
        painting_img = load_image_from_url(painting_url)
        result = apply_stretch(room_img, painting_img) if mode == "stretch" else apply_adapt(room_img, painting_img)

        return jsonify({"status": "ok", "mode": mode, "image_base64": image_to_base64(result)})

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to download image: {str(e)}"}), 400
    except ValueError as e:
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


@app.route("/noframe", methods=["POST"])
def noframe():
    try:
        data = request.get_json(force=True)
        painting_url = data.get("painting_url")
        if not painting_url:
            return jsonify({"error": "painting_url is required"}), 400

        painting_img = load_image_from_url(painting_url)
        result = apply_noframe(painting_img)

        return jsonify({"status": "ok", "image_base64": image_to_base64(result)})

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to download image: {str(e)}"}), 400
    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


@app.route("/zoom", methods=["POST"])
def zoom():
    try:
        data = request.get_json(force=True)
        painting_url = data.get("painting_url")
        mockup_url   = data.get("mockup_url")
        if not painting_url or not mockup_url:
            return jsonify({"error": "painting_url and mockup_url are required"}), 400

        painting_img = load_image_from_url(painting_url)
        mockup_img   = load_image_from_url(mockup_url)
        size_px      = int(data.get("size_px", 800))

        kwargs = {
            "frame_width":     data.get("frame_width"),
            "frame_color":     data.get("frame_color"),
            "shadow_offset_x": data.get("shadow_offset_x"),
            "shadow_offset_y": data.get("shadow_offset_y"),
            "shadow_blur":     data.get("shadow_blur"),
            "shadow_spread":   data.get("shadow_spread"),
            "shadow_opacity":  data.get("shadow_opacity"),
        }

        result = apply_zoom(painting_img, mockup_img, size_px, **kwargs)
        return jsonify({"status": "ok", "image_base64": image_to_base64(result)})

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to download image: {str(e)}"}), 400
    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


@app.route("/rect", methods=["POST"])
def rect():
    try:
        data = request.get_json(force=True)
        painting_url = data.get("painting_url")
        mockup_url   = data.get("mockup_url")
        if not painting_url or not mockup_url:
            return jsonify({"error": "painting_url and mockup_url are required"}), 400

        painting_img = load_image_from_url(painting_url)
        mockup_img   = load_image_from_url(mockup_url)
        size_px      = int(data.get("size_px", 800))

        kwargs = {
            "frame_width":     data.get("frame_width"),
            "frame_color":     data.get("frame_color"),
            "shadow_offset_x": data.get("shadow_offset_x"),
            "shadow_offset_y": data.get("shadow_offset_y"),
            "shadow_blur":     data.get("shadow_blur"),
            "shadow_spread":   data.get("shadow_spread"),
            "shadow_opacity":  data.get("shadow_opacity"),
        }

        result = apply_rect(painting_img, mockup_img, size_px, **kwargs)
        return jsonify({"status": "ok", "image_base64": image_to_base64(result)})

    except requests.exceptions.RequestException as e:
        return jsonify({"error": f"Failed to download image: {str(e)}"}), 400
    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


# ═════════════════════════════════════════════
# ENDPOINT: RECEIPT
# V23: Fire and forget — returns immediately, sends email in background
# ═════════════════════════════════════════════

@app.route("/receipt", methods=["POST"])
def receipt():
    try:
        data = request.get_json(force=True)

        # ── validate required fields ──
        required = ["customerEmail", "customerName", "orderNumber", "receiptNumber", "items", "total"]
        for field in required:
            if not data.get(field):
                return jsonify({"error": f"Missing required field: {field}"}), 400

        # ── build HTML ──
        html_body = build_receipt_html(data)

        # ── subject line ──
        doc_type      = data.get("documentType", "קבלה")
        receipt_num   = data.get("receiptNumber", "")
        business_name = data.get("businessName", "Artmidnet")
        subject = f"{doc_type} מספר {receipt_num} מאת {business_name}"

        # ── fire and forget — send in background ──
        to_email = data.get("customerEmail")
        thread = threading.Thread(
            target=send_receipt_email,
            args=(to_email, subject, html_body, data),
            daemon=True
        )
        thread.start()

        print(f"V28 /receipt: queued email to {to_email} | receipt={receipt_num} order={data.get('orderNumber')}")

        return jsonify({
            "status": "ok",
            "message": f"Receipt queued for {to_email}",
            "receiptNumber": receipt_num
        })

    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


# ═════════════════════════════════════════════
# ENDPOINTS: DOCX REPORTS
# ═════════════════════════════════════════════

@app.route("/layers-report", methods=["POST"])
def layers_report():
    try:
        data      = request.get_json(force=True)
        page_name = data.get("page_name", "Unknown Page")
        elements  = data.get("elements", [])
        if not elements:
            return jsonify({"error": "elements array is required"}), 400

        doc = DocxDocument()
        section = doc.sections[0]
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

        site_label = doc.add_paragraph()
        r0 = site_label.add_run('Artmidnet — Layers Report')
        r0.font.size = Pt(11); r0.font.color.rgb = RGBColor(0x88,0x88,0x88); r0.italic = True

        page_title = doc.add_paragraph()
        page_title.paragraph_format.space_before = Pt(4)
        page_title.paragraph_format.space_after  = Pt(2)
        r1 = page_title.add_run(f'Page: {page_name}')
        r1.bold = True; r1.font.size = Pt(26); r1.font.color.rgb = RGBColor(0x1F,0x38,0x64)
        pPr = page_title._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1F3864')
        pBdr.append(bottom); pPr.append(pBdr)

        meta = doc.add_paragraph()
        meta.paragraph_format.space_before = Pt(6)
        r2 = meta.add_run(f'Date: {datetime.date.today().strftime("%d/%m/%Y")}   |   Total elements: {len(elements)}')
        r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(0x88,0x88,0x88); r2.italic = True

        doc.add_paragraph()

        h2 = doc.add_paragraph()
        r = h2.add_run('Summary by Type')
        r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

        type_counts = {}
        for el in elements:
            t = el.get('type', 'Unknown')
            type_counts[t] = type_counts.get(t, 0) + 1

        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        hdr = summary_table.rows[0].cells
        for i, txt in enumerate(['Type', 'Count']):
            hdr[i].text = txt; set_cell_bg(hdr[i], '1F3864')
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(10)

        for t, count in sorted(type_counts.items(), key=lambda x: -x[1]):
            row = summary_table.add_row().cells
            color = get_type_color(t); rgb = hex_to_rgb(color)
            row[0].text = t; row[1].text = str(count)
            for cell in row:
                set_cell_bg(cell, 'F8F8F8')
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(*rgb); p.runs[0].font.size = Pt(10)

        doc.add_paragraph()
        h2b = doc.add_paragraph()
        r2b = h2b.add_run('Full Elements Tree')
        r2b.bold = True; r2b.font.size = Pt(13); r2b.font.color.rgb = RGBColor(0x2E,0x75,0xB6)

        note = doc.add_paragraph()
        rn = note.add_run('Depth column represents hierarchy level. ID is indented accordingly.')
        rn.font.size = Pt(9); rn.font.color.rgb = RGBColor(0x88,0x88,0x88); rn.italic = True
        doc.add_paragraph()

        main_table = doc.add_table(rows=1, cols=4)
        main_table.style = 'Table Grid'
        hdr2 = main_table.rows[0].cells
        for i, txt in enumerate(['ID', 'Type', 'Parent', 'Depth']):
            hdr2[i].text = txt; set_cell_bg(hdr2[i], '1F3864')
            run = hdr2[i].paragraphs[0].runs[0]
            run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(10)

        for el in elements:
            el_id = el.get('id',''); el_type = el.get('type','')
            el_parent = el.get('parent') or '—'; el_depth = int(el.get('depth', 1))
            color = get_type_color(el_type); rgb = hex_to_rgb(color)
            indent = '    ' * (el_depth - 1); row_bg = 'EBF3FB' if el_depth % 2 == 0 else 'FFFFFF'
            row = main_table.add_row().cells
            row[0].text = indent + '#' + el_id; row[1].text = el_type
            row[2].text = ('#' + el_parent) if el_parent != '—' else '—'; row[3].text = str(el_depth)
            for cell in row:
                set_cell_bg(cell, row_bg)
                p = cell.paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(9); p.runs[0].font.color.rgb = RGBColor(*rgb)

        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        docx_base64 = base64.b64encode(buf.read()).decode("utf-8")
        filename = f'Layers_Report_{page_name}_{datetime.date.today()}.docx'
        return jsonify({"status": "ok", "base64": docx_base64, "filename": filename})

    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


@app.route("/cms-report", methods=["POST"])
def cms_report():
    try:
        data        = request.get_json(force=True)
        collections = data.get("collections", [])

        if not collections:
            return jsonify({"error": "collections array is required"}), 400

        total_fields = sum(len(col.get('fields', [])) for col in collections)

        doc = DocxDocument()
        section = doc.sections[0]
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

        site_label = doc.add_paragraph()
        r0 = site_label.add_run('Artmidnet — CMS Schema Report')
        r0.bold = True; r0.font.size = Pt(20); r0.font.color.rgb = RGBColor(0x1F,0x38,0x64)

        meta = doc.add_paragraph()
        meta.paragraph_format.space_before = Pt(2)
        r_meta = meta.add_run(
            f'Collections: {len(collections)}  |  {total_fields} Fields  |  '
            f'Generated: {datetime.date.today().strftime("%d/%m/%Y")}'
        )
        r_meta.font.size = Pt(10); r_meta.font.color.rgb = RGBColor(0x88,0x88,0x88); r_meta.italic = True

        pPr = meta._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single'); bottom_bdr.set(qn('w:sz'), '6')
        bottom_bdr.set(qn('w:space'), '1'); bottom_bdr.set(qn('w:color'), '1F3864')
        pBdr.append(bottom_bdr); pPr.append(pBdr)

        doc.add_paragraph()

        for idx, col in enumerate(collections, 1):
            col_id   = col.get('collectionId', '')
            col_name = col.get('displayName', col_id)
            fields   = col.get('fields', [])

            col_heading = doc.add_paragraph()
            col_heading.paragraph_format.space_before = Pt(14)
            col_heading.paragraph_format.space_after  = Pt(2)
            r_col = col_heading.add_run(f'{idx}. Collection = {col_name}')
            r_col.bold = True; r_col.font.size = Pt(13); r_col.font.color.rgb = RGBColor(0x1F,0x38,0x64)

            col_sub = doc.add_paragraph()
            col_sub.paragraph_format.space_after = Pt(4)
            r_id = col_sub.add_run(f'id = {col_id}')
            r_id.font.size = Pt(10); r_id.font.color.rgb = RGBColor(0x88,0x88,0x88)

            col_count = doc.add_paragraph()
            col_count.paragraph_format.space_after = Pt(6)
            r_count = col_count.add_run(f'{len(fields)} fields')
            r_count.font.size = Pt(10); r_count.font.color.rgb = RGBColor(0x55,0x55,0x55)
            r_count.italic = True

            if fields:
                tbl = doc.add_table(rows=1, cols=4)
                tbl.style = 'Table Grid'

                hdr_cells = tbl.rows[0].cells
                for i, txt in enumerate(['Field Display Name', 'Field Type', 'Field Key', 'System']):
                    hdr_cells[i].text = txt
                    set_cell_bg(hdr_cells[i], '1F3864')
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                    run.font.size = Pt(10)

                for f_idx, field in enumerate(fields):
                    f_name   = field.get('displayName', '')
                    f_type   = field.get('type', '')
                    f_key    = field.get('key', '')
                    f_system = 'Yes' if field.get('systemField') else 'No'
                    row_bg   = 'F8F8F8' if f_idx % 2 == 0 else 'FFFFFF'
                    type_rgb = hex_to_rgb(get_field_type_color(f_type))

                    row = tbl.add_row().cells
                    row[0].text = f_name
                    row[1].text = f_type
                    row[2].text = f_key
                    row[3].text = f_system

                    for c_idx, cell in enumerate(row):
                        set_cell_bg(cell, row_bg)
                        p = cell.paragraphs[0]
                        if p.runs:
                            p.runs[0].font.size = Pt(10)
                            if c_idx == 1:
                                p.runs[0].font.color.rgb = RGBColor(*type_rgb)
                            else:
                                p.runs[0].font.color.rgb = RGBColor(0x33,0x33,0x33)

            doc.add_paragraph()

        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        docx_base64 = base64.b64encode(buf.read()).decode("utf-8")
        filename = f'CMS_Schema_Report_{datetime.date.today()}.docx'

        return jsonify({"status": "ok", "base64": docx_base64, "filename": filename})

    except Exception as e:
        return jsonify({"error": f"Internal error: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=False)
